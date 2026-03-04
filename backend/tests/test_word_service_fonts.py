"""Tests for WordExportService font_settings support."""
import io
import pytest
from unittest.mock import MagicMock
from docx.oxml.ns import qn

from backend.app.services.word_service import (
    WordExportService,
    DEFAULT_FONT_NAME,
    set_run_font,
    set_paragraph_font,
)
from backend.app.models.audit_schemas import FontSettings


def _make_outline_item(id_: str, title: str, content: str = "", children=None):
    """Create a mock outline item for testing."""
    item = MagicMock()
    item.id = id_
    item.title = title
    item.content = content
    item.children = children or []
    return item


class TestBuildDocumentDefaultFont:
    """build_document with font_settings=None uses DEFAULT_FONT_NAME."""

    def test_default_font_when_none(self):
        svc = WordExportService()
        items = [_make_outline_item("1", "章节一", "测试内容")]
        buf = svc.build_document(items, project_name="测试项目")
        assert isinstance(buf, io.BytesIO)
        assert svc._chinese_font == DEFAULT_FONT_NAME
        assert svc._english_font is None

    def test_backward_compatible_no_font_settings_arg(self):
        """Existing callers that don't pass font_settings still work."""
        svc = WordExportService()
        items = [_make_outline_item("1", "标题", "内容")]
        buf = svc.build_document(items, project_name="项目")
        assert buf.read(2) == b"PK"  # valid docx (zip) header


class TestBuildDocumentCustomFont:
    """build_document with FontSettings applies user fonts."""

    def test_chinese_font_applied(self):
        svc = WordExportService()
        fs = FontSettings(chinese_font="黑体")
        items = [_make_outline_item("1", "章节", "正文内容")]
        svc.build_document(items, project_name="测试", font_settings=fs)
        assert svc._chinese_font == "黑体"

    def test_english_font_applied(self):
        svc = WordExportService()
        fs = FontSettings(english_font="Arial")
        items = [_make_outline_item("1", "Section", "Body text")]
        svc.build_document(items, project_name="Test", font_settings=fs)
        assert svc._english_font == "Arial"

    def test_both_fonts_applied(self):
        svc = WordExportService()
        fs = FontSettings(chinese_font="仿宋", english_font="Calibri")
        items = [_make_outline_item("1", "章节", "内容")]
        svc.build_document(items, project_name="测试", font_settings=fs)
        assert svc._chinese_font == "仿宋"
        assert svc._english_font == "Calibri"

    def test_body_font_size_applied(self):
        svc = WordExportService()
        fs = FontSettings(body_font_size=12)
        items = [_make_outline_item("1", "章节", "内容")]
        svc.build_document(items, project_name="测试", font_settings=fs)
        assert svc._body_font_size == 12

    def test_document_styles_use_custom_chinese_font(self):
        svc = WordExportService()
        fs = FontSettings(chinese_font="楷体")
        items = [_make_outline_item("1", "章节", "内容")]
        svc.build_document(items, project_name="测试", font_settings=fs)
        # Verify the Normal style has the custom font in EastAsia
        normal_style = svc.doc.styles["Normal"]
        rpr = normal_style._element.rPr
        east_asia = rpr.rFonts.get(qn("w:eastAsia"))
        assert east_asia == "楷体"

    def test_document_styles_use_custom_english_font(self):
        svc = WordExportService()
        fs = FontSettings(chinese_font="宋体", english_font="Georgia")
        items = [_make_outline_item("1", "章节", "内容")]
        svc.build_document(items, project_name="测试", font_settings=fs)
        normal_style = svc.doc.styles["Normal"]
        assert normal_style.font.name == "Georgia"

    def test_output_is_valid_docx(self):
        svc = WordExportService()
        fs = FontSettings(chinese_font="黑体", english_font="Arial", body_font_size=11)
        items = [_make_outline_item("1", "章节一", "测试内容")]
        buf = svc.build_document(items, project_name="测试项目", font_settings=fs)
        data = buf.read()
        assert data[:2] == b"PK"  # valid zip/docx


class TestModuleLevelFunctions:
    """Module-level set_run_font / set_paragraph_font still work independently."""

    def test_set_run_font_default(self):
        import docx as d
        doc = d.Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_run_font(run)
        assert run.font.name == DEFAULT_FONT_NAME

    def test_set_run_font_custom(self):
        import docx as d
        doc = d.Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_run_font(run, "黑体")
        assert run.font.name == "黑体"

    def test_set_paragraph_font(self):
        import docx as d
        doc = d.Document()
        p = doc.add_paragraph("hello")
        set_paragraph_font(p, "楷体")
        for run in p.runs:
            assert run.font.name == "楷体"
