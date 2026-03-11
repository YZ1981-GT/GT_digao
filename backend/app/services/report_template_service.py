"""报告模板服务。

管理致同审计报告模板（国企版/上市版），模板存储为结构化 Markdown。
支持知识库持久化、内存缓存、章节级检索、Word 导入。
"""
import logging
import re
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from ..models.audit_schemas import (
    ReportTemplateDocument,
    ReportTemplateSection,
    ReportTemplateType,
    TemplateCategory,
    TemplateTocEntry,
)

logger = logging.getLogger(__name__)

# 缓存 key: (template_type, template_category)
_TemplateKey = Tuple[str, str]
_SectionKey = Tuple[str, str, str]  # (type, category, path)


class ReportTemplateService:
    """报告模板服务。"""

    def __init__(self):
        self._cache: Dict[_TemplateKey, ReportTemplateDocument] = {}
        self._section_cache: Dict[_SectionKey, str] = {}

    # ─── Public API ───

    def get_template(
        self,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
    ) -> Optional[ReportTemplateDocument]:
        """获取指定模板内容。"""
        key = (template_type.value, template_category.value)
        if key in self._cache:
            return self._cache[key]

        # 尝试从知识库加载
        doc = self._load_from_storage(template_type, template_category)
        if doc:
            self._cache[key] = doc
        return doc

    def get_template_section(
        self,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
        section_path: str,
    ) -> Optional[str]:
        """按层级路径精确检索章节内容。"""
        sec_key = (template_type.value, template_category.value, section_path)
        if sec_key in self._section_cache:
            return self._section_cache[sec_key]

        doc = self.get_template(template_type, template_category)
        if not doc:
            return None

        for section in doc.sections:
            if section.path == section_path:
                self._section_cache[sec_key] = section.content
                return section.content
        return None

    def get_template_toc(
        self,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
    ) -> List[TemplateTocEntry]:
        """获取模板目录结构（仅标题层级）。"""
        doc = self.get_template(template_type, template_category)
        if not doc:
            return []

        entries = []
        for section in doc.sections:
            has_children = any(
                s.path.startswith(section.path + "/") for s in doc.sections
                if s.path != section.path
            )
            entries.append(TemplateTocEntry(
                path=section.path,
                level=section.level,
                title=section.title,
                has_children=has_children,
            ))
        return entries

    def update_template(
        self,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
        content: str,
    ) -> ReportTemplateDocument:
        """更新模板内容，自动失效缓存。"""
        if not content.strip():
            raise ValueError("模板内容不能为空")

        sections = self._parse_markdown_sections(content)
        now = datetime.now().isoformat()
        doc = ReportTemplateDocument(
            template_type=template_type,
            template_category=template_category,
            full_content=content,
            sections=sections,
            version=now,
            updated_at=now,
        )

        # 失效缓存
        key = (template_type.value, template_category.value)
        self._cache[key] = doc
        # 清除该模板的章节缓存
        to_remove = [k for k in self._section_cache if k[0] == template_type.value and k[1] == template_category.value]
        for k in to_remove:
            del self._section_cache[k]

        # 持久化
        self._save_to_storage(template_type, template_category, content)
        return doc

    def import_from_word(
        self,
        file_content: bytes,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
    ) -> ReportTemplateDocument:
        """从 Word 文档导入并转换为结构化 Markdown。"""
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(file_content))
        except Exception as e:
            raise ValueError(f"Word 文件转换失败: {e}")

        md_lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style:
                md_lines.append(f"# {text}")
            elif "heading 2" in style:
                md_lines.append(f"## {text}")
            elif "heading 3" in style:
                md_lines.append(f"### {text}")
            else:
                md_lines.append(text)
            md_lines.append("")

        content = "\n".join(md_lines)
        if not content.strip():
            raise ValueError("Word 文件内容为空")

        return self.update_template(template_type, template_category, content)

    def clear_cache(self):
        """清除所有缓存。"""
        self._cache.clear()
        self._section_cache.clear()

    # ─── Markdown 解析 ───

    @staticmethod
    def _parse_markdown_sections(content: str) -> List[ReportTemplateSection]:
        """解析 Markdown 为层级章节列表。"""
        sections: List[ReportTemplateSection] = []
        current_path_parts: List[str] = []
        current_content_lines: List[str] = []
        current_title = ""
        current_level = 0

        def _flush():
            if current_title:
                path = "/".join(current_path_parts) if current_path_parts else current_title
                sections.append(ReportTemplateSection(
                    path=path,
                    level=current_level,
                    title=current_title,
                    content="\n".join(current_content_lines).strip(),
                ))

        for line in content.split("\n"):
            heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
            if heading_match:
                _flush()
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                current_level = level
                current_title = title
                current_content_lines = []

                # 更新路径
                while len(current_path_parts) >= level:
                    current_path_parts.pop()
                current_path_parts.append(title)
            else:
                current_content_lines.append(line)

        _flush()
        return sections

    # ─── 存储 ───

    def _load_from_storage(
        self,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
    ) -> Optional[ReportTemplateDocument]:
        """从知识库加载模板。

        优先按精确 filename 匹配（report_template_soe_report_body 等），
        若未找到则按文件名关键字（国企/上市 + 附注等）自动聚合原始上传文档。
        """
        try:
            from .knowledge_service import knowledge_service
            exact_filename = f"report_template_{template_type.value}_{template_category.value}"
            docs = knowledge_service.get_documents("report_templates")

            # 1) 精确匹配
            for doc in docs:
                if doc.get("filename") == exact_filename:
                    content = knowledge_service.get_document_content("report_templates", doc["id"])
                    if content and content.strip():
                        sections = self._parse_markdown_sections(content)
                        return ReportTemplateDocument(
                            template_type=template_type,
                            template_category=template_category,
                            full_content=content,
                            sections=sections,
                            version=doc.get("created_at", datetime.now().isoformat()),
                            updated_at=doc.get("created_at", datetime.now().isoformat()),
                        )

            # 2) 按文件名关键字聚合原始上传文档
            type_keywords = {
                ReportTemplateType.SOE: ["国企", "soe"],
                ReportTemplateType.LISTED: ["上市", "listed"],
            }
            cat_keywords_include = {
                TemplateCategory.NOTES: ["附注", "notes"],
                TemplateCategory.REPORT_BODY: [],  # 默认分类，排除附注即可
            }
            cat_keywords_exclude = {
                TemplateCategory.REPORT_BODY: ["附注", "notes"],
                TemplateCategory.NOTES: [],
            }

            tk = type_keywords.get(template_type, [])
            ci = cat_keywords_include.get(template_category, [])
            ce = cat_keywords_exclude.get(template_category, [])

            matched_parts: list[str] = []
            for doc in docs:
                fn = doc.get("filename", "")
                fn_lower = fn.lower()
                # 跳过精确命名的合并文档
                if fn.startswith("report_template_"):
                    continue
                # 匹配模板类型
                if not any(kw in fn or kw in fn_lower for kw in tk):
                    continue
                # 匹配分类（附注 vs 正文）
                if ci and not any(kw in fn or kw in fn_lower for kw in ci):
                    continue
                if ce and any(kw in fn or kw in fn_lower for kw in ce):
                    continue
                content = knowledge_service.get_document_content("report_templates", doc["id"])
                if content and content.strip():
                    matched_parts.append(f"# {fn}\n\n{content}")

            if matched_parts:
                merged = "\n\n".join(matched_parts)
                sections = self._parse_markdown_sections(merged)
                now = datetime.now().isoformat()
                return ReportTemplateDocument(
                    template_type=template_type,
                    template_category=template_category,
                    full_content=merged,
                    sections=sections,
                    version=now,
                    updated_at=now,
                )
        except Exception as e:
            logger.debug("从知识库加载模板失败: %s", e)
        return None

    def _save_to_storage(
        self,
        template_type: ReportTemplateType,
        template_category: TemplateCategory,
        content: str,
    ):
        """持久化模板到知识库（覆盖已有同名文档）。"""
        try:
            from .knowledge_service import knowledge_service
            filename = f"report_template_{template_type.value}_{template_category.value}"
            # 删除已有同名文档，避免重复
            docs = knowledge_service.get_documents("report_templates")
            for doc in docs:
                if doc.get("filename") == filename:
                    knowledge_service.delete_document("report_templates", doc["id"])
            knowledge_service.add_document("report_templates", filename, content)
        except Exception as e:
            logger.warning("保存模板到知识库失败: %s", e)


# 模块级单例
report_template_service = ReportTemplateService()
