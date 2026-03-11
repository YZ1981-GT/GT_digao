"""审计底稿文件解析服务。

复用现有 FileService 的文件提取能力（extract_text_from_pdf, extract_text_from_docx），
新增 Excel 解析（openpyxl/xlrd）和底稿编号识别能力。
"""
import asyncio
import functools
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from ..models.audit_schemas import (
    CellData,
    ExcelParseResult,
    PdfParseResult,
    SheetData,
    WordParseResult,
    WorkpaperClassification,
    WorkpaperParseResult,
    WorkpaperType,
)

logger = logging.getLogger(__name__)

# Optional imports — graceful degradation
try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    logger.warning("openpyxl 未安装，xlsx 解析不可用")

try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False
    logger.warning("xlrd 未安装，xls 解析不可用")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx 未安装，docx 解析不可用")


class WorkpaperParser:
    """审计底稿文件解析服务。"""

    SUPPORTED_FORMATS = {'.xlsx', '.xls', '.doc', '.docx', '.pdf'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    # 底稿编号正则：B-xx, C-xx, D-xx ~ M-xx, Q-xx
    WORKPAPER_ID_PATTERN = r'^([B-M])-?\d+'

    # 业务循环映射
    BUSINESS_CYCLE_MAP = {
        'D': '销售循环',
        'E': '货币资金循环',
        'F': '存货循环',
        'G': '投资循环',
        'H': '固定资产循环',
        'I': '无形资产循环',
        'J': '职工薪酬循环',
        'K': '管理循环',
        'L': '债务循环',
        'M': '权益循环',
        'Q': '关联方循环',
    }


    # ─── Public API ───

    async def parse_file(self, file_path: str, filename: str) -> WorkpaperParseResult:
        """统一入口：校验文件大小和格式，分发到对应解析器。"""
        now_iso = datetime.now(timezone.utc).isoformat()
        result_id = str(uuid.uuid4())
        ext = os.path.splitext(filename)[1].lower()

        # 基础校验
        if not os.path.exists(file_path):
            return self._error_result(result_id, filename, ext, 0, now_iso, "文件不存在")

        file_size = os.path.getsize(file_path)

        if file_size == 0:
            return self._error_result(result_id, filename, ext, file_size, now_iso, "文件为空")

        if file_size > self.MAX_FILE_SIZE:
            return self._error_result(
                result_id, filename, ext, file_size, now_iso,
                f"文件大小超过限制（{self.MAX_FILE_SIZE // (1024 * 1024)}MB）"
            )

        if ext not in self.SUPPORTED_FORMATS:
            return self._error_result(
                result_id, filename, ext, file_size, now_iso,
                f"不支持的文件格式：{ext}，支持格式：{', '.join(sorted(self.SUPPORTED_FORMATS))}"
            )

        # 分发解析
        try:
            content_text = ""
            structured_data: Optional[Dict[str, Any]] = None

            if ext in ('.xlsx', '.xls'):
                excel_result = await self.parse_excel(file_path, ext)
                structured_data = excel_result.model_dump()
                # 将 Excel 内容拼接为纯文本
                content_text = self._excel_to_text(excel_result)

            elif ext in ('.docx', '.doc'):
                word_result = await self.parse_word(file_path)
                structured_data = word_result.model_dump()
                content_text = self._word_to_text(word_result)

            elif ext == '.pdf':
                pdf_result = await self.parse_pdf(file_path)
                structured_data = pdf_result.model_dump()
                content_text = pdf_result.text

            classification = self.identify_workpaper_type(filename, content_text)

            return WorkpaperParseResult(
                id=result_id,
                filename=filename,
                file_format=ext.lstrip('.'),
                file_size=file_size,
                classification=classification,
                content_text=content_text,
                structured_data=structured_data,
                parse_status="success",
                parsed_at=now_iso,
            )

        except Exception as e:
            logger.exception("解析文件 %s 失败", filename)
            error_msg = str(e)
            # 友好化常见错误
            if "encrypted" in error_msg.lower() or "password" in error_msg.lower():
                error_msg = "文件已加密，无法解析"
            elif "corrupt" in error_msg.lower() or "invalid" in error_msg.lower():
                error_msg = f"文件损坏或格式无效：{error_msg}"
            return self._error_result(result_id, filename, ext, file_size, now_iso, error_msg)

    async def parse_excel(self, file_path: str, ext: str = "") -> ExcelParseResult:
        """解析 Excel 文件，提取工作表、单元格数据、公式和合并单元格信息。"""
        if not ext:
            ext = os.path.splitext(file_path)[1].lower()

        if ext == '.xlsx':
            return await asyncio.to_thread(self._parse_xlsx, file_path)
        elif ext == '.xls':
            return await asyncio.to_thread(self._parse_xls, file_path)
        else:
            raise ValueError(f"不支持的 Excel 格式：{ext}")

    async def parse_word(self, file_path: str) -> WordParseResult:
        """解析 Word 文件，提取段落、表格、标题层级和批注。支持 .docx 和 .doc 格式。"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.doc':
            return await self._parse_doc_legacy(file_path)

        # .docx 使用 python-docx — 在线程池中执行以避免阻塞事件循环
        return await asyncio.to_thread(self._parse_docx_sync, file_path)

    # ── 从 docx_to_md 移植的增强标题检测逻辑 ──

    # 财务报表附注中常见的描述性标题行模式
    _NOTE_TITLE_PATTERNS = re.compile('|'.join([
        r'^按.{1,20}(计提|分类|归集|划分|披露)',
        r'^组合计提项目[：:]',
        r'^期末本公司',
        r'^本期计提',
        r'^转回或收回',
        r'^本期实际核销',
        r'^重要的.{2,30}$',
        r'^按单项计提',
        r'^按组合计提',
        r'^按坏账计提',
        r'^按(预付|欠款方)归集',
        r'^作为(承租人|出租人)',
        r'^外币货币性项目$',
        r'^境外经营实体$',
    ]))

    @staticmethod
    def _detect_heading_level_enhanced(para, style_name: str, text: str,
                                        next_is_table: bool = False) -> Optional[int]:
        """增强的标题级别检测（融合 docx_to_md 的多种检测策略）。

        检测顺序：
        1. Word 内置 Heading / 标题 样式
        2. 自定义样式模糊匹配（如"标题3的样式"）
        3. outlineLvl XML 属性
        4. 加粗短段落 → 子标题
        5. 财务报表附注模式匹配标题
        6. 表格前描述性短段落 → 子标题
        """
        if not text:
            return None

        # 1. Word 内置样式
        if style_name.startswith("Heading"):
            try:
                return int(style_name.replace("Heading", "").strip())
            except ValueError:
                pass
        if style_name.startswith("标题"):
            try:
                return int(style_name.replace("标题", "").strip())
            except ValueError:
                pass

        # 2. 自定义样式模糊匹配
        if style_name:
            m = re.search(r'标题\s*([1-9])', style_name)
            if m:
                return int(m.group(1))

        # 3. outlineLvl XML 属性
        from docx.oxml.ns import qn
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                val = outlineLvl.get(qn('w:val'))
                if val is not None:
                    lvl = int(val) + 1
                    if 1 <= lvl <= 6:
                        return lvl

        # 4. 加粗短段落 → 子标题
        runs = [r for r in para.runs if r.text.strip()]
        is_all_bold = runs and all(r.bold for r in runs)
        if is_all_bold and len(text) <= 80 and text[-1] not in '。；;':
            # 不把括号包裹的内容当标题
            if not ((text.startswith('（') and text.endswith('）')) or
                    (text.startswith('(') and text.endswith(')'))):
                return None  # 标记为 bold_title，由调用方处理

        # 返回 None 表示非标题
        return None

    def _parse_docx_sync(self, file_path: str) -> WordParseResult:
        """同步解析 .docx 文件（在线程池中调用）。
        融合 docx_to_md.py 的增强标题检测能力。"""
        if not HAS_DOCX:
            raise RuntimeError("python-docx 未安装，无法解析 docx 文件")

        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise RuntimeError(f"Word 文件打开失败：{e}")

        paragraphs: List[Dict[str, Any]] = []
        headings: List[Dict[str, Any]] = []
        tables: List[List[List[str]]] = []
        comments: List[Dict[str, str]] = []

        # ── 按文档 body 元素顺序统一遍历段落和表格 ──
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph

        # 先收集所有 body 元素，以便前瞻判断（如"下一个是否为表格"）
        body_elements = []
        for element in doc.element.body:
            tag = element.tag
            if tag == qn('w:p'):
                body_elements.append(('para', DocxParagraph(element, doc)))
            elif tag == qn('w:tbl'):
                body_elements.append(('table', DocxTable(element, doc)))

        table_after_para_idx: List[int] = []
        para_idx = -1
        current_heading_level = 0

        for idx, (etype, obj) in enumerate(body_elements):
            if etype == 'para':
                para = obj
                style_name = para.style.name if para.style else ""
                text = para.text.strip()

                # 前瞻：下一个非空元素是否为表格
                next_is_table = False
                for j in range(idx + 1, min(idx + 4, len(body_elements))):
                    ntype, nobj = body_elements[j]
                    if ntype == 'table':
                        next_is_table = True
                        break
                    if ntype == 'para' and nobj.text.strip():
                        break

                level = self._detect_heading_level_enhanced(para, style_name, text, next_is_table)

                # 补充检测：加粗短段落 / 模式匹配标题 / 表格前描述行
                if level is None and text:
                    runs = [r for r in para.runs if r.text.strip()]
                    is_all_bold = runs and all(r.bold for r in runs)
                    is_short = len(text) <= 80 and text[-1] not in '。；;'
                    not_wrapped = not ((text.startswith('（') and text.endswith('）')) or
                                       (text.startswith('(') and text.endswith(')')))

                    if is_all_bold and is_short and not_wrapped:
                        # 加粗短段落 → 当前标题级别 + 1
                        level = min(current_heading_level + 1, 6) if current_heading_level else 4

                    elif self._NOTE_TITLE_PATTERNS.search(text) and is_short and not_wrapped:
                        # 财务报表附注模式匹配标题
                        level = min(current_heading_level + 1, 6) if current_heading_level else 4

                    elif (next_is_table and len(text) <= 60 and
                          text[-1] not in '。；;.' and not_wrapped and
                          not re.match(r'^续[（(：:]', text) and text != '续' and
                          not re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)):
                        # 表格前描述性短段落 → 子标题
                        level = min(current_heading_level + 1, 6) if current_heading_level else 4

                para_info: Dict[str, Any] = {
                    "text": para.text,
                    "style": style_name,
                }
                if level is not None:
                    para_info["level"] = level
                    headings.append({"text": para.text, "level": level})
                    current_heading_level = level

                paragraphs.append(para_info)
                para_idx = len(paragraphs) - 1

            elif etype == 'table':
                tbl = obj
                table_data: List[List[str]] = []
                for row in tbl.rows:
                    # 去重合并单元格
                    seen_cells: set = set()
                    row_data: List[str] = []
                    for cell in row.cells:
                        cell_id = id(cell._tc)
                        if cell_id in seen_cells:
                            continue
                        seen_cells.add(cell_id)
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
                table_after_para_idx.append(para_idx)

        # 兼容旧字段 table_contexts
        table_contexts: List[str] = []
        for tpi in table_after_para_idx:
            if 0 <= tpi < len(paragraphs):
                table_contexts.append(paragraphs[tpi].get('text', '').strip())
            else:
                table_contexts.append("")

        # 批注（python-docx 不直接支持批注 API，通过 XML 解析）
        try:
            from lxml import etree
            comments_part = None
            for rel in doc.part.rels.values():
                if "comments" in rel.reltype:
                    comments_part = rel.target_part
                    break
            if comments_part is not None:
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                root = etree.fromstring(comments_part.blob)
                for comment_el in root.findall('.//w:comment', ns):
                    author = comment_el.get(f'{{{ns["w"]}}}author', '')
                    texts = comment_el.findall('.//w:t', ns)
                    text = ''.join(t.text or '' for t in texts)
                    if text:
                        comments.append({"author": author, "text": text})
        except Exception:
            # 批注提取失败不影响主流程
            pass

        return WordParseResult(
            paragraphs=paragraphs,
            tables=tables,
            headings=headings,
            comments=comments,
            table_contexts=table_contexts,
            table_after_para_idx=table_after_para_idx,
        )

    async def _parse_doc_legacy(self, file_path: str) -> WordParseResult:
        """解析旧版 .doc 文件。Windows 上使用 pywin32 COM 接口。"""
        return await asyncio.to_thread(self._parse_doc_legacy_sync, file_path)

    def _parse_doc_legacy_sync(self, file_path: str) -> WordParseResult:
        """同步解析 .doc 文件（在线程池中调用）。"""
        text = ""
        tables: List[List[List[str]]] = []
        headings: List[Dict[str, Any]] = []
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            try:
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = False
                abs_path = os.path.abspath(file_path)
                doc = word_app.Documents.Open(abs_path, ReadOnly=True)
                text = doc.Content.Text or ""
                # 提取标题（通过 OutlineLevel 判断）
                for para in doc.Paragraphs:
                    try:
                        outline_level = para.OutlineLevel
                        # OutlineLevel: 1-9 为标题级别，10 (wdOutlineLevelBodyText) 为正文
                        if 1 <= outline_level <= 9:
                            para_text = para.Range.Text.strip().rstrip('\r\x07')
                            if para_text:
                                headings.append({"text": para_text, "level": outline_level})
                    except Exception:
                        pass
                # 提取表格
                for table in doc.Tables:
                    table_data: List[List[str]] = []
                    for row_idx in range(1, table.Rows.Count + 1):
                        row_data: List[str] = []
                        for col_idx in range(1, table.Columns.Count + 1):
                            try:
                                cell_text = table.Cell(row_idx, col_idx).Range.Text
                                cell_text = cell_text.strip().rstrip('\r\x07')
                                row_data.append(cell_text)
                            except Exception:
                                row_data.append("")
                        table_data.append(row_data)
                    tables.append(table_data)
                doc.Close(False)
                word_app.Quit()
            finally:
                pythoncom.CoUninitialize()
        except ImportError:
            raise RuntimeError("pywin32 未安装，Windows 上无法解析 .doc 文件。请安装：pip install pywin32")
        except Exception as e:
            raise RuntimeError(f".doc 文件解析失败：{e}")

        # 按换行拆分为段落
        paragraphs = [{"text": line, "style": ""} for line in text.split('\r') if line.strip()]

        return WordParseResult(
            paragraphs=paragraphs,
            tables=tables,
            headings=headings,
            comments=[],
        )

    async def parse_pdf(self, file_path: str) -> PdfParseResult:
        """解析 PDF 文件，复用 FileService 的提取逻辑。"""
        from .file_service import FileService

        text = await FileService.extract_text_from_pdf(file_path)

        # 提取表格（从文本中解析 [表格 N] ... [表格结束] 块）
        tables = self._extract_tables_from_text(text)

        # 页数
        page_count = self._count_pdf_pages(file_path)

        return PdfParseResult(
            text=text,
            tables=tables,
            page_count=page_count,
        )

    def identify_workpaper_type(self, filename: str, content: str) -> WorkpaperClassification:
        """识别底稿编号体系（B/C/D-M/Q 类）和业务循环分类。"""
        # 先从文件名匹配
        basename = os.path.splitext(os.path.basename(filename))[0]
        match = re.match(self.WORKPAPER_ID_PATTERN, basename)

        # 文件名未匹配则尝试从内容前 500 字符匹配
        if not match and content:
            match = re.search(self.WORKPAPER_ID_PATTERN, content[:500], re.MULTILINE)

        if not match:
            return WorkpaperClassification()

        letter = match.group(1).upper()
        workpaper_id = match.group(0)

        # 映射到 WorkpaperType 枚举
        wp_type: Optional[WorkpaperType] = None
        try:
            wp_type = WorkpaperType(letter)
        except ValueError:
            pass

        business_cycle = self.BUSINESS_CYCLE_MAP.get(letter)

        return WorkpaperClassification(
            workpaper_type=wp_type,
            business_cycle=business_cycle,
            workpaper_id=workpaper_id,
        )

    async def batch_parse(self, files: List[Tuple[str, str]]) -> List[WorkpaperParseResult]:
        """批量解析多个底稿文件，并发执行以提升性能。

        Args:
            files: [(file_path, filename), ...]
        """
        tasks = [self.parse_file(fp, fn) for fp, fn in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        # 将异常转为错误结果
        final: List[WorkpaperParseResult] = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                fp, fn = files[i]
                logger.error("批量解析文件 %s 失败: %s", fn, result)
                final.append(self._error_result(
                    str(uuid.uuid4()), fn,
                    os.path.splitext(fn)[1].lower(),
                    0, datetime.now(timezone.utc).isoformat(),
                    str(result),
                ))
            else:
                final.append(result)
        return final


    # ─── Private helpers ───

    def _parse_xlsx(self, file_path: str) -> ExcelParseResult:
        """使用 openpyxl 解析 .xlsx 文件。"""
        if not HAS_OPENPYXL:
            raise RuntimeError("openpyxl 未安装，无法解析 xlsx 文件")

        try:
            wb = openpyxl.load_workbook(file_path, data_only=False, read_only=False)
        except Exception as e:
            raise RuntimeError(f"Excel 文件打开失败：{e}")

        sheets: List[SheetData] = []
        sheet_names: List[str] = list(wb.sheetnames)

        for ws in wb.worksheets:
            cells: List[CellData] = []
            merged_ranges = [str(mr) for mr in ws.merged_cells.ranges]

            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None or cell.data_type == 'f':
                        formula = None
                        value = cell.value
                        if isinstance(value, str) and value.startswith('='):
                            formula = value
                        cells.append(CellData(
                            row=cell.row,
                            col=cell.column,
                            value=self._safe_cell_value(value),
                            formula=formula,
                            is_merged=self._is_in_merged(cell.row, cell.column, ws.merged_cells.ranges),
                        ))

            sheets.append(SheetData(
                name=ws.title,
                cells=cells,
                merged_ranges=merged_ranges,
            ))

        wb.close()
        return ExcelParseResult(sheets=sheets, sheet_names=sheet_names)

    def _parse_xls(self, file_path: str) -> ExcelParseResult:
        """使用 xlrd 解析 .xls 文件。"""
        if not HAS_XLRD:
            raise RuntimeError("xlrd 未安装，无法解析 xls 文件")

        try:
            wb = xlrd.open_workbook(file_path, formatting_info=False)
        except Exception as e:
            raise RuntimeError(f"Excel 文件打开失败：{e}")

        sheets: List[SheetData] = []
        sheet_names: List[str] = wb.sheet_names()

        for ws in wb.sheets():
            cells: List[CellData] = []
            merged_ranges: List[str] = []

            # xlrd merged cells: list of (row_lo, row_hi, col_lo, col_hi)
            for rlo, rhi, clo, chi in ws.merged_cells:
                merged_ranges.append(f"{self._col_letter(clo+1)}{rlo+1}:{self._col_letter(chi)}{rhi}")

            merged_set = set()
            for rlo, rhi, clo, chi in ws.merged_cells:
                for r in range(rlo, rhi):
                    for c in range(clo, chi):
                        merged_set.add((r, c))

            for row_idx in range(ws.nrows):
                for col_idx in range(ws.ncols):
                    value = ws.cell_value(row_idx, col_idx)
                    if value != '':
                        cells.append(CellData(
                            row=row_idx + 1,
                            col=col_idx + 1,
                            value=self._safe_cell_value(value),
                            formula=None,  # xlrd 不支持公式提取
                            is_merged=(row_idx, col_idx) in merged_set,
                        ))

            sheets.append(SheetData(
                name=ws.name,
                cells=cells,
                merged_ranges=merged_ranges,
            ))

        return ExcelParseResult(sheets=sheets, sheet_names=sheet_names)

    @staticmethod
    def _is_in_merged(row: int, col: int, merged_ranges) -> bool:
        """检查单元格是否在合并区域内。"""
        for mr in merged_ranges:
            if (mr.min_row <= row <= mr.max_row and mr.min_col <= col <= mr.max_col):
                return True
        return False

    @staticmethod
    def _col_letter(col_num: int) -> str:
        """将列号转为字母（1→A, 2→B, ...）。"""
        result = ""
        while col_num > 0:
            col_num, remainder = divmod(col_num - 1, 26)
            result = chr(65 + remainder) + result
        return result

    @staticmethod
    def _safe_cell_value(value: Any) -> Any:
        """将单元格值转为 JSON 安全类型。"""
        if value is None:
            return None
        if isinstance(value, (int, float, str, bool)):
            return value
        if isinstance(value, datetime):
            return value.isoformat()
        return str(value)

    @staticmethod
    def _excel_to_text(result: ExcelParseResult) -> str:
        """将 Excel 解析结果拼接为纯文本。"""
        parts: List[str] = []
        for sheet in result.sheets:
            parts.append(f"--- 工作表: {sheet.name} ---")
            # 按行组织
            rows_map: Dict[int, List[Tuple[int, Any]]] = {}
            for cell in sheet.cells:
                rows_map.setdefault(cell.row, []).append((cell.col, cell.value))
            for row_num in sorted(rows_map.keys()):
                cols = sorted(rows_map[row_num], key=lambda x: x[0])
                row_text = " | ".join(str(v) if v is not None else "" for _, v in cols)
                parts.append(row_text)
        return "\n".join(parts)

    @staticmethod
    def _word_to_text(result: WordParseResult) -> str:
        """将 Word 解析结果拼接为纯文本。"""
        parts: List[str] = []
        for para in result.paragraphs:
            text = para.get("text", "")
            if text:
                parts.append(text)
        for i, table in enumerate(result.tables, 1):
            parts.append(f"\n[表格 {i}]")
            for row in table:
                parts.append(" | ".join(row))
            parts.append("[表格结束]")
        return "\n".join(parts)

    @staticmethod
    def _extract_tables_from_text(text: str) -> List[List[List[str]]]:
        """从 FileService 提取的文本中解析表格块。"""
        tables: List[List[List[str]]] = []
        pattern = re.compile(r'\[表格\s*\d+\](.*?)\[表格结束\]', re.DOTALL)
        for match in pattern.finditer(text):
            block = match.group(1).strip()
            table: List[List[str]] = []
            for line in block.split('\n'):
                line = line.strip()
                if line and '|' in line:
                    row = [cell.strip() for cell in line.split('|')]
                    table.append(row)
            if table:
                tables.append(table)
        return tables

    @staticmethod
    def _count_pdf_pages(file_path: str) -> int:
        """获取 PDF 页数。"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return len(reader.pages)
        except Exception:
            return 0

    @staticmethod
    def _error_result(
        result_id: str,
        filename: str,
        ext: str,
        file_size: int,
        parsed_at: str,
        error_message: str,
    ) -> WorkpaperParseResult:
        """构造错误解析结果。"""
        return WorkpaperParseResult(
            id=result_id,
            filename=filename,
            file_format=ext.lstrip('.'),
            file_size=file_size,
            classification=WorkpaperClassification(),
            content_text="",
            structured_data=None,
            parse_status="error",
            error_message=error_message,
            parsed_at=parsed_at,
        )
