"""Word文档导出服务"""
import io
import os
import re
import logging
from typing import List, Optional

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from ..models.audit_schemas import FontSettings

logger = logging.getLogger(__name__)

# 默认字体配置（可通过环境变量覆盖）
DEFAULT_FONT_NAME = os.environ.get("WORD_EXPORT_FONT", "宋体")


def set_run_font(run: docx.text.run.Run, font_name: str = DEFAULT_FONT_NAME) -> None:
    """统一将 run 字体设置为指定字体（包含 EastAsia 字体设置）"""
    run.font.name = font_name
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph: docx.text.paragraph.Paragraph, font_name: str = DEFAULT_FONT_NAME) -> None:
    """将段落内所有 runs 字体设置为指定字体"""
    for run in paragraph.runs:
        set_run_font(run, font_name)


class WordExportService:
    """Word文档导出服务"""

    def __init__(self):
        self.doc = docx.Document()
        self._chinese_font: str = DEFAULT_FONT_NAME
        self._english_font: Optional[str] = None
        self._body_font_size: Optional[float] = None
        self._setup_styles()

    def _apply_font_settings(self, font_settings: Optional[FontSettings] = None) -> None:
        """Apply user-specified font settings, updating internal font state and document styles."""
        if font_settings:
            if font_settings.chinese_font:
                self._chinese_font = font_settings.chinese_font
            if font_settings.english_font:
                self._english_font = font_settings.english_font
            if font_settings.body_font_size:
                self._body_font_size = font_settings.body_font_size
        self._setup_styles()

    def _set_run_font(self, run: docx.text.run.Run) -> None:
        """Apply the resolved chinese/english fonts to a single run."""
        set_run_font(run, self._chinese_font)
        if self._english_font:
            run.font.name = self._english_font
            # Keep EastAsia as chinese font
            r = run._element.rPr
            if r is not None and r.rFonts is not None:
                r.rFonts.set(qn("w:eastAsia"), self._chinese_font)

    def _set_paragraph_font(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Apply the resolved fonts to all runs in a paragraph."""
        for run in paragraph.runs:
            self._set_run_font(run)

    def _setup_styles(self) -> None:
        """统一设置文档的基础字体"""
        try:
            styles = self.doc.styles
            base_styles = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]
            for style_name in base_styles:
                if style_name in styles:
                    style = styles[style_name]
                    font = style.font
                    font.name = self._english_font or self._chinese_font
                    if style._element.rPr is None:
                        style._element._add_rPr()
                    rpr = style._element.rPr
                    rpr.rFonts.set(qn("w:eastAsia"), self._chinese_font)
                    if style_name == "Normal":
                        font.bold = False
                        if self._body_font_size:
                            font.size = Pt(self._body_font_size)
        except Exception:
            pass

    def _add_markdown_runs(self, para: docx.text.paragraph.Paragraph, text: str) -> None:
        """在指定段落中追加 markdown 文本的 runs"""
        pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
        parts = re.split(pattern, text)
        for part in parts:
            if not part:
                continue
            run = para.add_run()
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                run.text = part[1:-1]
            else:
                run.text = part
            self._set_run_font(run)

    def _add_markdown_paragraph(self, text: str) -> None:
        """将一段 Markdown 文本解析为一个普通段落"""
        para = self.doc.add_paragraph()
        self._add_markdown_runs(para, text)
        para.paragraph_format.space_after = Pt(6)

    def _parse_markdown_blocks(self, content: str):
        """识别 Markdown 内容中的块级元素，返回结构化的 block 列表"""
        blocks = []
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip("\r").strip()
            if not line:
                i += 1
                continue

            # 列表项（有序/无序）
            if line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
                items = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if stripped.startswith("- ") or stripped.startswith("* "):
                        text = re.sub(r"^[-*]\s+", "", stripped).strip()
                        if text:
                            items.append(("unordered", None, text))
                        i += 1
                        continue
                    m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                    if m_num:
                        num_str, text = m_num.groups()
                        text = text.strip()
                        if text:
                            items.append(("ordered", num_str, text))
                        i += 1
                        continue
                    break
                if items:
                    blocks.append(("list", items))
                continue

            # 表格
            if "|" in line:
                rows = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if "|" in stripped:
                        if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                            cells = [c.strip() for c in stripped.split("|")]
                            row_text = " | ".join([c for c in cells if c])
                            if row_text:
                                rows.append(row_text)
                        i += 1
                    else:
                        break
                if rows:
                    blocks.append(("table", rows))
                continue

            # Markdown 标题
            if line.startswith("#"):
                m = re.match(r"^(#+)\s*(.*)$", line)
                if m:
                    level_marks, title_text = m.groups()
                    level = min(len(level_marks), 3)
                    blocks.append(("heading", level, title_text.strip()))
                i += 1
                continue

            # 普通段落
            para_lines = []
            while i < len(lines):
                raw = lines[i].rstrip("\r")
                stripped = raw.strip()
                if (
                    stripped
                    and not stripped.startswith("-")
                    and not stripped.startswith("*")
                    and "|" not in stripped
                    and not stripped.startswith("#")
                ):
                    para_lines.append(stripped)
                    i += 1
                else:
                    break
            if para_lines:
                text = " ".join(para_lines)
                blocks.append(("paragraph", text))
            else:
                i += 1

        return blocks

    def _render_markdown_blocks(self, blocks) -> None:
        """将结构化的 Markdown blocks 渲染到文档"""
        for block in blocks:
            kind = block[0]
            if kind == "list":
                items = block[1]
                for item_kind, num_str, text in items:
                    p = self.doc.add_paragraph()
                    if item_kind == "unordered":
                        run = p.add_run("• ")
                        self._set_run_font(run)
                    else:
                        prefix = f"{num_str}."
                        run = p.add_run(prefix + " ")
                        self._set_run_font(run)
                    self._add_markdown_runs(p, text)
            elif kind == "table":
                rows = block[1]
                for row in rows:
                    self._add_markdown_paragraph(row)
            elif kind == "heading":
                _, level, text = block
                heading = self.doc.add_heading(text, level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._set_paragraph_font(heading)
            elif kind == "paragraph":
                _, text = block
                self._add_markdown_paragraph(text)

    def _add_markdown_content(self, content: str) -> None:
        """解析并渲染 Markdown 文本到文档"""
        blocks = self._parse_markdown_blocks(content)
        self._render_markdown_blocks(blocks)

    def _add_outline_items(self, items, level: int = 1) -> None:
        """递归构建文档内容（章节和内容）"""
        for item in items:
            # 章节标题
            if level <= 3:
                heading = self.doc.add_heading(f"{item.id} {item.title}", level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for hr in heading.runs:
                    self._set_run_font(hr)
            else:
                para = self.doc.add_paragraph()
                run = para.add_run(f"{item.id} {item.title}")
                run.bold = True
                self._set_run_font(run)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(3)

            # 叶子节点内容
            if not item.children:
                content = item.content or ""
                if content.strip():
                    self._add_markdown_content(content)
            else:
                self._add_outline_items(item.children, level + 1)

    def build_document(
        self,
        outline_items,
        project_name: Optional[str] = None,
        project_overview: Optional[str] = None,
        font_settings: Optional[FontSettings] = None,
    ) -> io.BytesIO:
        """
        构建完整的 Word 文档并返回内存中的字节流。

        Args:
            outline_items: 目录结构列表（Pydantic OutlineItem 对象）
            project_name: 项目名称
            project_overview: 项目概述
            font_settings: 用户自定义字体设置，None 时使用 DEFAULT_FONT_NAME

        Returns:
            io.BytesIO: 可直接用于 StreamingResponse 的字节流
        """
        # Apply user font settings if provided
        if font_settings:
            self._apply_font_settings(font_settings)

        # AI 生成声明
        p = self.doc.add_paragraph()
        run = p.add_run("内容由AI生成")
        run.italic = True
        run.font.size = Pt(9)
        self._set_run_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 文档标题
        title = project_name or "审计文档"
        title_p = self.doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        self._set_run_font(title_run)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 项目概述
        if project_overview:
            heading = self.doc.add_heading("项目概述", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._set_paragraph_font(heading)
            overview_p = self.doc.add_paragraph(project_overview)
            self._set_paragraph_font(overview_p)
            overview_p.paragraph_format.space_after = Pt(12)

        # 递归构建章节内容
        self._add_outline_items(outline_items)

        # 输出到内存
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
