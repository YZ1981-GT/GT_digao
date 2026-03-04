"""复核报告生成与导出服务。

将 ReviewEngine 的复核结果整合为结构化复核报告，
支持导出为 Word（.docx）和 PDF 格式。
Word 导出复用现有 word_service.py 的文档生成能力，
PDF 导出使用 weasyprint 将 HTML 渲染为 PDF。
"""
import io
import uuid
import logging
from datetime import datetime
from typing import List, Optional, Dict

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..models.audit_schemas import (
    ReviewReport,
    ReviewFinding,
    RiskLevel,
)
from .word_service import (
    set_run_font,
    set_paragraph_font,
    DEFAULT_FONT_NAME,
)

logger = logging.getLogger(__name__)

# 风险等级排序权重（高→中→低）
_RISK_SORT_ORDER = {
    RiskLevel.HIGH: 0,
    RiskLevel.MEDIUM: 1,
    RiskLevel.LOW: 2,
}

# 风险等级中文标签
_RISK_LABELS = {
    RiskLevel.HIGH: "高风险",
    RiskLevel.MEDIUM: "中风险",
    RiskLevel.LOW: "低风险",
}


class ReportGenerator:
    """复核报告生成与导出服务。"""

    def __init__(self):
        pass  # No external service dependencies needed for report generation

    # ── 报告生成 ──

    def generate_report(
        self,
        workpaper_ids: List[str],
        dimensions: List[str],
        findings: List[ReviewFinding],
        conclusion: str,
        project_id: Optional[str] = None,
    ) -> ReviewReport:
        """将复核结果整合为结构化复核报告。

        报告结构：复核概要、问题清单（按风险等级排序）、修改建议、复核结论。
        自动计算风险等级统计汇总。
        """
        # 按风险等级排序：HIGH → MEDIUM → LOW
        sorted_findings = sorted(
            findings,
            key=lambda f: _RISK_SORT_ORDER.get(f.risk_level, 99),
        )

        # 统计各风险等级数量
        summary: Dict[str, int] = {"high": 0, "medium": 0, "low": 0}
        for f in sorted_findings:
            level_key = f.risk_level.value  # "high" / "medium" / "low"
            summary[level_key] = summary.get(level_key, 0) + 1

        return ReviewReport(
            id=str(uuid.uuid4()),
            workpaper_ids=workpaper_ids,
            dimensions=dimensions,
            findings=sorted_findings,
            summary=summary,
            conclusion=conclusion,
            reviewed_at=datetime.now().isoformat(),
            project_id=project_id,
        )

    # ── Word 导出 ──

    def export_to_word(self, report: ReviewReport) -> bytes:
        """导出复核报告为 Word 格式（.docx）。

        报告结构：标题 → 概要表格 → 按风险等级分组的问题清单 → 结论。
        复用 word_service.py 的字体设置工具函数。
        """
        doc = docx.Document()

        # ── 标题 ──
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("审计底稿复核报告")
        title_run.bold = True
        title_run.font.size = Pt(18)
        set_run_font(title_run, DEFAULT_FONT_NAME)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── 概要信息表格 ──
        heading_summary = doc.add_heading("一、复核概要", level=1)
        set_paragraph_font(heading_summary, DEFAULT_FONT_NAME)

        summary_table = doc.add_table(rows=4, cols=2, style="Table Grid")
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        summary_rows = [
            ("复核时间", report.reviewed_at),
            ("复核维度", "、".join(report.dimensions)),
            ("底稿数量", str(len(report.workpaper_ids))),
            (
                "风险统计",
                f"高风险 {report.summary.get('high', 0)} 项 / "
                f"中风险 {report.summary.get('medium', 0)} 项 / "
                f"低风险 {report.summary.get('low', 0)} 项",
            ),
        ]
        for i, (label, value) in enumerate(summary_rows):
            cell_label = summary_table.cell(i, 0)
            cell_value = summary_table.cell(i, 1)
            cell_label.text = label
            cell_value.text = value
            for cell in (cell_label, cell_value):
                for para in cell.paragraphs:
                    set_paragraph_font(para, DEFAULT_FONT_NAME)

        # ── 问题清单（按风险等级分组） ──
        heading_findings = doc.add_heading("二、问题清单", level=1)
        set_paragraph_font(heading_findings, DEFAULT_FONT_NAME)

        for risk_level in (RiskLevel.HIGH, RiskLevel.MEDIUM, RiskLevel.LOW):
            level_findings = [
                f for f in report.findings if f.risk_level == risk_level
            ]
            if not level_findings:
                continue

            level_heading = doc.add_heading(
                f"{_RISK_LABELS[risk_level]}（{len(level_findings)} 项）",
                level=2,
            )
            set_paragraph_font(level_heading, DEFAULT_FONT_NAME)

            for idx, finding in enumerate(level_findings, 1):
                # 问题标题
                p_title = doc.add_paragraph()
                run_title = p_title.add_run(f"{idx}. [{finding.dimension}] {finding.description}")
                run_title.bold = True
                set_run_font(run_title, DEFAULT_FONT_NAME)

                # 问题详情
                details = [
                    ("位置", finding.location),
                    ("参考依据", finding.reference),
                    ("修改建议", finding.suggestion),
                    ("状态", finding.status.value),
                ]
                for label, value in details:
                    if value:
                        p_detail = doc.add_paragraph()
                        run_label = p_detail.add_run(f"  {label}：")
                        run_label.bold = True
                        set_run_font(run_label, DEFAULT_FONT_NAME)
                        run_value = p_detail.add_run(value)
                        set_run_font(run_value, DEFAULT_FONT_NAME)

        # ── 结论 ──
        heading_conclusion = doc.add_heading("三、复核结论", level=1)
        set_paragraph_font(heading_conclusion, DEFAULT_FONT_NAME)

        if report.conclusion:
            p_conclusion = doc.add_paragraph(report.conclusion)
            set_paragraph_font(p_conclusion, DEFAULT_FONT_NAME)

        # 输出字节
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    # ── PDF 导出 ──

    def export_to_pdf(self, report: ReviewReport) -> bytes:
        """导出复核报告为 PDF 格式。

        使用 weasyprint 将 HTML 渲染为 PDF。
        A4 尺寸（210mm × 297mm），页边距上下 20mm、左右 15mm，黑白配色。
        """
        html_content = self._build_report_html(report)
        from weasyprint import HTML  # lazy import to avoid startup cost

        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes

    def _build_report_html(self, report: ReviewReport) -> str:
        """构建复核报告的 HTML 模板（A4 黑白配色）。"""
        # 按风险等级分组构建 findings HTML
        findings_html = ""
        for risk_level in (RiskLevel.HIGH, RiskLevel.MEDIUM, RiskLevel.LOW):
            level_findings = [
                f for f in report.findings if f.risk_level == risk_level
            ]
            if not level_findings:
                continue

            findings_html += f'<h2>{_RISK_LABELS[risk_level]}（{len(level_findings)} 项）</h2>\n'
            findings_html += '<table>\n<thead><tr>'
            findings_html += '<th>序号</th><th>维度</th><th>描述</th><th>位置</th><th>建议</th><th>状态</th>'
            findings_html += '</tr></thead>\n<tbody>\n'

            for idx, finding in enumerate(level_findings, 1):
                findings_html += "<tr>"
                findings_html += f"<td>{idx}</td>"
                findings_html += f"<td>{_escape_html(finding.dimension)}</td>"
                findings_html += f"<td>{_escape_html(finding.description)}</td>"
                findings_html += f"<td>{_escape_html(finding.location)}</td>"
                findings_html += f"<td>{_escape_html(finding.suggestion)}</td>"
                findings_html += f"<td>{_escape_html(finding.status.value)}</td>"
                findings_html += "</tr>\n"

            findings_html += "</tbody>\n</table>\n"

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<style>
@page {{
    size: A4;
    margin: 20mm 15mm;
}}
body {{
    font-family: "SimSun", "宋体", serif;
    font-size: 12pt;
    color: #000;
    line-height: 1.6;
}}
h1 {{
    text-align: center;
    font-size: 18pt;
    margin-bottom: 20px;
}}
h2 {{
    font-size: 14pt;
    border-bottom: 1px solid #000;
    padding-bottom: 4px;
    margin-top: 20px;
}}
h3 {{
    font-size: 12pt;
    margin-top: 16px;
}}
table {{
    width: 100%;
    border-collapse: collapse;
    margin: 10px 0;
    page-break-inside: auto;
}}
th, td {{
    border: 1px solid #000;
    padding: 6px 8px;
    text-align: left;
    font-size: 10pt;
    word-wrap: break-word;
}}
th {{
    background-color: #eee;
    font-weight: bold;
}}
tr {{
    page-break-inside: avoid;
}}
.summary-table td:first-child {{
    width: 120px;
    font-weight: bold;
}}
.conclusion {{
    margin-top: 20px;
    padding: 10px;
    border: 1px solid #000;
}}
</style>
</head>
<body>
<h1>审计底稿复核报告</h1>

<h2>一、复核概要</h2>
<table class="summary-table">
<tr><td>复核时间</td><td>{_escape_html(report.reviewed_at)}</td></tr>
<tr><td>复核维度</td><td>{_escape_html("、".join(report.dimensions))}</td></tr>
<tr><td>底稿数量</td><td>{len(report.workpaper_ids)}</td></tr>
<tr><td>风险统计</td><td>高风险 {report.summary.get("high", 0)} 项 / 中风险 {report.summary.get("medium", 0)} 项 / 低风险 {report.summary.get("low", 0)} 项</td></tr>
</table>

<h2>二、问题清单</h2>
{findings_html}

<h2>三、复核结论</h2>
<div class="conclusion">{_escape_html(report.conclusion)}</div>
</body>
</html>"""
        return html

    # ── 序列化 / 反序列化 ──

    def parse_report_to_structured(self, report: ReviewReport) -> dict:
        """将复核报告解析为结构化数据格式（字典）。"""
        return report.model_dump(mode="json")

    def structured_to_report(self, data: dict) -> ReviewReport:
        """从结构化数据重建复核报告对象（往返一致性）。"""
        return ReviewReport.model_validate(data)


def _escape_html(text: str) -> str:
    """简单 HTML 转义。"""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
