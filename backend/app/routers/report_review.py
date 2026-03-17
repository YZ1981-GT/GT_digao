"""审计报告复核路由层。

注册到 /api/report-review/ 前缀。
"""
import asyncio
import json
import logging
import os
import uuid
from datetime import datetime
from typing import Dict, List, Optional
from urllib.parse import quote as requests_quote

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from ..models.audit_schemas import (
    FindingConfirmationStatus,
    FindingConversation,
    FindingStatus,
    MatchingMap,
    ReportFileType,
    ReportReviewConfig,
    ReportReviewFinding,
    ReportReviewFindingCategory,
    ReportReviewResult,
    ReportReviewSession,
    ReportTemplateType,
    RiskLevel,
    TemplateCategory,
)
from ..services.report_review_engine import report_review_engine
from ..services.report_parser import report_parser
from ..services.report_template_service import report_template_service
from ..services.reconciliation_engine import reconciliation_engine
from ..services.table_structure_analyzer import TableStructureAnalyzer
from ..services.file_service import FileService
from ..services import session_store

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/report-review", tags=["report-review"])

# 内存会话存储（启动时从磁盘恢复）
_sessions: Dict[str, ReportReviewSession] = {}
_findings: Dict[str, List[ReportReviewFinding]] = {}
_conversations: Dict[str, FindingConversation] = {}

# 恢复持久化数据
_restored = session_store.restore_all()
_sessions.update(_restored[0])
_findings.update(_restored[1])
_conversations.update(_restored[2])
del _restored


# ─── Request/Response Models ───

class UploadRequest(BaseModel):
    template_type: str = "soe"

class ConfirmMatchingRequest(BaseModel):
    session_id: str
    matching_map: MatchingMap

class StartReviewRequest(BaseModel):
    session_id: str
    template_type: str = "soe"
    prompt_id: Optional[str] = None
    custom_prompt: Optional[str] = None
    change_threshold: float = 0.3
    change_amount_threshold: float = 0

class EditFindingRequest(BaseModel):
    description: Optional[str] = None
    suggestion: Optional[str] = None
    risk_level: Optional[str] = None

class ChatRequest(BaseModel):
    message: str

class TraceRequest(BaseModel):
    trace_type: str  # cross_reference / template_compare / data_drill_down

class BatchRequest(BaseModel):
    finding_ids: List[str]
    action: str  # confirm / dismiss


class AddAnnotationRequest(BaseModel):
    session_id: str
    section_title: str
    account_name: str = ""
    description: str
    risk_level: str = "medium"

class StatusUpdateRequest(BaseModel):
    status: str  # open / resolved


# ─── 文件上传与解析 ───

@router.post("/upload")
async def upload_files(
    files: List[UploadFile] = File(...),
    template_type: str = Form("soe"),
):
    """上传审计报告文件并解析。"""
    import tempfile
    import aiofiles

    try:
        tt = ReportTemplateType(template_type)
    except ValueError:
        raise HTTPException(400, f"不支持的模板类型: {template_type}")

    session_id = str(uuid.uuid4())[:8]
    session = ReportReviewSession(
        id=session_id,
        template_type=tt,
        created_at=datetime.now().isoformat(),
    )

    # 保存文件到临时目录并解析
    temp_files: List[str] = []
    try:
        for file in files:
            if file.size and file.size > 50 * 1024 * 1024:
                raise HTTPException(413, f"文件 {file.filename} 超过50MB限制")

            filename = file.filename or "unknown"
            ext = os.path.splitext(filename)[1].lower()
            content = await file.read()

            # 保存到临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            temp_files.append(tmp_path)

            file_id = str(uuid.uuid4())[:8]
            session.file_ids.append(file_id)

            # 记录文件名映射
            session.source_file_names[file_id] = filename

            # 生成页面截图（PDF 文件）
            if ext == '.pdf':
                try:
                    page_dir = os.path.join("uploads", "pages", session_id, file_id)
                    page_map = FileService.render_pdf_page_images(tmp_path, page_dir, dpi=150)
                    if page_map:
                        if not session.page_image_dir:
                            session.page_image_dir = os.path.join("uploads", "pages", session_id)
                        logger.info("PDF 页面截图生成: %s, %d 页", filename, len(page_map))
                except Exception as e:
                    logger.warning("页面截图生成失败 %s: %s", filename, e)

            # 分类文件
            try:
                if ext in ('.doc', '.docx', '.xlsx', '.xls', '.pdf'):
                    content_text = ""
                else:
                    try:
                        content_text = content.decode('utf-8')[:2000]
                    except (UnicodeDecodeError, AttributeError):
                        content_text = content.decode('utf-8', errors='ignore')[:2000]
                file_type = report_parser.classify_report_file(filename, content_text)
                session.file_classifications[file_id] = file_type
            except Exception as e:
                raise HTTPException(400, f"文件分类失败: {e}")

            # 实际解析文件内容（根据文件分类决定提取方式）
            try:
                if ext in ('.xlsx', '.xls'):
                    excel_result = await report_parser.parse_excel(tmp_path, ext)
                    sheets = report_parser.extract_sheets(excel_result)
                    session.sheet_data[file_id] = sheets
                    for sheet in sheets:
                        items = report_parser.extract_statement_items(sheet)
                        session.statement_items.extend(items)

                elif ext in ('.docx', '.doc'):
                    word_result = await report_parser.parse_word(tmp_path)
                    # 用解析后的文本重新分类（上面分类时 content_text 为空）
                    word_text = ' '.join(
                        p.get('text', '') for p in word_result.paragraphs[:50]
                    )
                    file_type = report_parser.classify_report_file(filename, word_text)
                    session.file_classifications[file_id] = file_type

                    if file_type == ReportFileType.NOTES_TO_STATEMENTS:
                        # 附注文件：提取表格和层级结构
                        note_tables = report_parser.extract_note_tables(word_result)
                        session.note_tables.extend(note_tables)
                        note_sections = report_parser.extract_note_sections(word_result, note_tables)
                        session.note_sections.extend(note_sections)
                    elif file_type == ReportFileType.AUDIT_REPORT_BODY:
                        # 审计报告正文：提取段落内容
                        for para in word_result.paragraphs:
                            text = para.get('text', '').strip()
                            if not text:
                                # 保留空行用于段落分隔
                                session.audit_report_content.append({
                                    'text': '', 'level': None, 'style': '',
                                })
                                continue
                            session.audit_report_content.append({
                                'text': text,
                                'level': para.get('level'),
                                'style': para.get('style', ''),
                            })
                    else:
                        logger.info(f"Word 文件 {filename} 分类为 {file_type.value}")

                elif ext == '.pdf':
                    pdf_result = await report_parser.parse_pdf(tmp_path)
                    logger.info(f"PDF 文件 {filename} 分类为 {file_type.value}")

            except Exception as e:
                logger.warning(f"文件 {filename} 解析失败: {e}")

        session.status = "parsed"

        # 自动构建科目匹配映射（基于标准模板）
        if session.statement_items and session.note_tables:
            try:
                session.matching_map = reconciliation_engine.build_matching_map(
                    session.statement_items, session.note_tables
                )
                logger.info("自动匹配完成: %d 条映射, %d 未匹配",
                            len(session.matching_map.entries),
                            len(session.matching_map.unmatched_items))
            except Exception as e:
                logger.warning("自动匹配失败: %s", e)

            # 预先用规则识别表格结构（毫秒级），复核阶段直接复用
            try:
                analyzer = TableStructureAnalyzer()
                for note in session.note_tables:
                    ts = analyzer._analyze_with_rules(note)
                    session.table_structures[note.id] = ts
                logger.info("表格结构预识别完成: %d 个表格", len(session.table_structures))
            except Exception as e:
                logger.warning("表格结构预识别失败: %s", e)

        _sessions[session_id] = session
        session_store.save_session(session_id, session)
        return {
            "session_id": session_id,
            "file_count": len(files),
            "template_type": template_type,
            "statement_items": len(session.statement_items),
            "note_tables": len(session.note_tables),
            "note_sections": len(session.note_sections),
        }
    finally:
        # 清理临时文件
        for tmp_path in temp_files:
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass


@router.post("/parse")
async def parse_session(session_id: str = Form(...)):
    """解析已上传文件。"""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "会话不存在")

    session.status = "parsed"
    _sessions[session_id] = session
    session_store.save_session(session_id, session)
    return {"status": "parsed", "session_id": session_id}


@router.get("/session/{session_id}")
async def get_session(session_id: str):
    """获取会话状态。"""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "会话不存在")
    return session.model_dump()


@router.get("/session/{session_id}/sheets")
async def get_sheets(session_id: str):
    """获取已解析的 Sheet 列表。"""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "会话不存在")
    return {"sheets": session.sheet_data}


@router.get("/session/{session_id}/note-tables")
async def get_note_tables(session_id: str, account_name: Optional[str] = None, note_id: Optional[str] = None, note_ids: Optional[str] = None):
    """获取附注表格数据，可按科目名、表格ID或多个ID筛选。"""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "会话不存在")
    tables = session.note_tables

    # 优先按 note_ids（逗号分隔）精确查找
    if note_ids:
        id_list = [x.strip() for x in note_ids.split(",") if x.strip()]
        id_set = set(id_list)
        tables = [t for t in tables if t.id in id_set]
    elif note_id:
        tables = [t for t in tables if t.id == note_id]
    elif account_name:
        # 多策略匹配：精确 → 包含 → 模糊
        exact = [t for t in tables if t.account_name == account_name]
        if exact:
            tables = exact
        else:
            contains = [t for t in tables if account_name in t.account_name
                        or t.account_name in account_name
                        or account_name in t.section_title
                        or t.account_name in account_name.replace("按", "").replace("分类披露", "").replace("计提方法", "")]
            tables = contains if contains else []
    return {"note_tables": [t.model_dump() for t in tables]}


@router.get("/session/{session_id}/page-image/{file_id}/{page_num}")
async def get_page_image(session_id: str, file_id: str, page_num: int):
    """获取源文档指定页的截图。"""
    from fastapi.responses import FileResponse

    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "会话不存在")

    img_path = os.path.join("uploads", "pages", session_id, file_id, f"page_{page_num}.jpg")
    if not os.path.isfile(img_path):
        raise HTTPException(404, "页面截图不存在")

    return FileResponse(img_path, media_type="image/jpeg")


@router.get("/session/{session_id}/page-images-info")
async def get_page_images_info(session_id: str):
    """获取会话中所有可用的页面截图信息。"""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "会话不存在")

    result = {}
    base_dir = os.path.join("uploads", "pages", session_id)
    if os.path.isdir(base_dir):
        for fid in os.listdir(base_dir):
            fid_dir = os.path.join(base_dir, fid)
            if os.path.isdir(fid_dir):
                pages = sorted([
                    int(f.replace("page_", "").replace(".jpg", ""))
                    for f in os.listdir(fid_dir) if f.startswith("page_") and f.endswith(".jpg")
                ])
                result[fid] = {
                    "filename": session.source_file_names.get(fid, fid),
                    "pages": pages,
                }
    return {"page_images": result}


# ─── 匹配与复核 ───

@router.post("/confirm-matching")
async def confirm_matching(req: ConfirmMatchingRequest):
    """用户确认/调整科目匹配映射。"""
    session = _sessions.get(req.session_id)
    if not session:
        raise HTTPException(404, "会话不存在")
    session.matching_map = req.matching_map
    session.status = "matched"
    return {"status": "matched"}


@router.post("/start")
async def start_review(req: StartReviewRequest):
    """发起复核，SSE 流式返回。"""
    session = _sessions.get(req.session_id)
    if not session:
        raise HTTPException(404, "会话不存在")

    try:
        tt = ReportTemplateType(req.template_type)
    except ValueError:
        tt = session.template_type

    config = ReportReviewConfig(
        session_id=req.session_id,
        template_type=tt,
        prompt_id=req.prompt_id,
        custom_prompt=req.custom_prompt,
        change_threshold=req.change_threshold,
        change_amount_threshold=req.change_amount_threshold,
    )

    async def event_stream():
        async for event in report_review_engine.review_stream(session, config):
            yield f"data: {event}\n\n"
            await asyncio.sleep(0)  # 强制 flush，确保 SSE 事件立即发送到客户端
            # 收集 findings
            try:
                data = json.loads(event)
                if data.get("status") == "completed" and "result" in data:
                    result = data["result"]
                    if "findings" in result:
                        new_findings = [
                            ReportReviewFinding(**f) for f in result["findings"]
                        ]
                        # 保留已有的手动批注（manual_annotation），与新复核结果合并
                        existing = _findings.get(req.session_id, [])
                        annotations = [
                            f for f in existing
                            if f.category == ReportReviewFindingCategory.MANUAL_ANNOTATION
                        ]
                        _findings[req.session_id] = annotations + new_findings
                        session_store.save_findings(req.session_id, _findings[req.session_id])
            except Exception:
                pass

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─── Finding 交互 ───

@router.get("/findings/{session_id}")
async def get_findings(session_id: str):
    """获取所有 Finding。"""
    findings = _findings.get(session_id, [])
    return {"findings": [f.model_dump() for f in findings]}


@router.post("/finding/annotation")
async def add_annotation(req: AddAnnotationRequest):
    """手动插入批注（复核意见）。"""
    import uuid
    finding = ReportReviewFinding(
        id=str(uuid.uuid4())[:8],
        category=ReportReviewFindingCategory.MANUAL_ANNOTATION,
        risk_level=RiskLevel(req.risk_level),
        account_name=req.account_name or req.section_title,
        location=f"附注-{req.section_title}",
        description=req.description,
        suggestion="",
        confirmation_status=FindingConfirmationStatus.PENDING_CONFIRMATION,
    )
    _findings.setdefault(req.session_id, []).append(finding)
    session_store.save_findings(req.session_id, _findings[req.session_id])
    return {"status": "created", "finding": finding.model_dump()}


@router.get("/finding/{finding_id}")
async def get_finding(finding_id: str):
    """获取单条 Finding 详情。"""
    for session_findings in _findings.values():
        for f in session_findings:
            if f.id == finding_id:
                conv = _conversations.get(finding_id)
                return {
                    "finding": f.model_dump(),
                    "conversation": conv.model_dump() if conv else None,
                }
    raise HTTPException(404, "Finding 不存在")


@router.patch("/finding/{finding_id}/edit")
async def edit_finding(finding_id: str, req: EditFindingRequest):
    """编辑 Finding。"""
    for session_findings in _findings.values():
        for f in session_findings:
            if f.id == finding_id:
                if req.description is not None:
                    f.description = req.description
                if req.suggestion is not None:
                    f.suggestion = req.suggestion
                if req.risk_level is not None:
                    f.risk_level = RiskLevel(req.risk_level)
                _persist_finding_change(finding_id)
                return {"status": "updated"}
    raise HTTPException(404, "Finding 不存在")


@router.patch("/finding/{finding_id}/confirm")
async def confirm_finding(finding_id: str):
    """确认 Finding。"""
    for session_findings in _findings.values():
        for f in session_findings:
            if f.id == finding_id:
                f.confirmation_status = FindingConfirmationStatus.CONFIRMED
                _persist_finding_change(finding_id)
                return {"status": "confirmed"}
    raise HTTPException(404, "Finding 不存在")


@router.patch("/finding/{finding_id}/dismiss")
async def dismiss_finding(finding_id: str):
    """忽略 Finding。"""
    for session_findings in _findings.values():
        for f in session_findings:
            if f.id == finding_id:
                f.confirmation_status = FindingConfirmationStatus.DISMISSED
                _persist_finding_change(finding_id)
                return {"status": "dismissed"}
    raise HTTPException(404, "Finding 不存在")


@router.patch("/finding/{finding_id}/restore")
async def restore_finding(finding_id: str):
    """恢复已忽略的 Finding。"""
    for session_findings in _findings.values():
        for f in session_findings:
            if f.id == finding_id:
                f.confirmation_status = FindingConfirmationStatus.PENDING_CONFIRMATION
                _persist_finding_change(finding_id)
                return {"status": "restored"}
    raise HTTPException(404, "Finding 不存在")


@router.post("/findings/batch")
async def batch_findings(req: BatchRequest):
    """批量确认/忽略。"""
    status_map = {
        "confirm": FindingConfirmationStatus.CONFIRMED,
        "dismiss": FindingConfirmationStatus.DISMISSED,
    }
    new_status = status_map.get(req.action)
    if not new_status:
        raise HTTPException(400, f"不支持的操作: {req.action}")

    updated = 0
    affected_sessions: set = set()
    for sid, session_findings in _findings.items():
        for f in session_findings:
            if f.id in req.finding_ids:
                f.confirmation_status = new_status
                updated += 1
                affected_sessions.add(sid)
    for sid in affected_sessions:
        session_store.save_findings(sid, _findings[sid])
    return {"updated": updated}


# ─── Finding 对话与溯源 ───

@router.post("/finding/{finding_id}/chat")
async def chat_finding(finding_id: str, req: ChatRequest):
    """用户追问，SSE 流式返回。"""
    finding = _find_finding(finding_id)
    if not finding:
        raise HTTPException(404, "Finding 不存在")

    conv = _conversations.setdefault(finding_id, FindingConversation(finding_id=finding_id))

    async def event_stream():
        async for event in report_review_engine.chat_about_finding(
            finding, req.message, conv, report_review_engine.openai_service
        ):
            yield f"data: {event}\n\n"
        # 对话结束后持久化
        sid = _find_finding_session_id(finding_id)
        if sid:
            session_store.save_conversations(sid, {k: v for k, v in _conversations.items() if _find_finding_session_id(k) == sid})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/finding/{finding_id}/trace")
async def trace_finding(finding_id: str, req: TraceRequest):
    """溯源分析，SSE 流式返回。"""
    finding = _find_finding(finding_id)
    if not finding:
        raise HTTPException(404, "Finding 不存在")

    conv = _conversations.setdefault(finding_id, FindingConversation(finding_id=finding_id))

    async def event_stream():
        async for event in report_review_engine.trace_finding(
            finding, req.trace_type, conv, report_review_engine.openai_service
        ):
            yield f"data: {event}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.get("/finding/{finding_id}/conversation")
async def get_conversation(finding_id: str):
    """获取 Finding 对话历史。"""
    conv = _conversations.get(finding_id)
    if not conv:
        return {"messages": []}
    return conv.model_dump()


# ─── 报告与导出 ───

@router.get("/result/{session_id}")
async def get_result(session_id: str):
    """获取复核结果。如果用户未勾选任何问题，则导出全部问题。"""
    findings = _findings.get(session_id, [])
    confirmed = [f for f in findings if f.confirmation_status == FindingConfirmationStatus.CONFIRMED]
    dismissed = [f for f in findings if f.confirmation_status == FindingConfirmationStatus.DISMISSED]
    pending = [f for f in findings if f.confirmation_status == FindingConfirmationStatus.PENDING_CONFIRMATION]

    # 如果没有任何已确认的问题，则输出所有未忽略的问题（pending + confirmed）
    output = confirmed if confirmed else [f for f in findings if f.confirmation_status != FindingConfirmationStatus.DISMISSED]

    category_summary = {}
    risk_summary = {"high": 0, "medium": 0, "low": 0}
    for f in output:
        cat = f.category.value
        category_summary[cat] = category_summary.get(cat, 0) + 1
        risk_summary[f.risk_level.value] += 1

    # 生成结论
    total = len(output)
    if total == 0:
        conclusion = "本次复核未发现明显问题。"
    else:
        high = sum(1 for f in output if f.risk_level.value == "high")
        medium = sum(1 for f in output if f.risk_level.value == "medium")
        low = sum(1 for f in output if f.risk_level.value == "low")
        parts = [f"本次复核共发现 {total} 个待确认问题"]
        details = []
        if high:
            details.append(f"高风险 {high} 个")
        if medium:
            details.append(f"中风险 {medium} 个")
        if low:
            details.append(f"低风险 {low} 个")
        if details:
            parts.append(f"（{'、'.join(details)}）")
        parts.append("。所有问题需经用户确认后纳入最终报告。")
        conclusion = "".join(parts)

    return {
        "session_id": session_id,
        "findings": [f.model_dump() for f in output],
        "category_summary": category_summary,
        "risk_summary": risk_summary,
        "confirmation_summary": {
            "confirmed": len(confirmed),
            "dismissed": len(dismissed),
            "pending": len(pending),
        },
        "conclusion": conclusion,
        "total_findings": len(findings),
    }


class ExportRequest(BaseModel):
    session_id: str
    format: str = "word"


@router.post("/export")
async def export_report(req: ExportRequest):
    """导出复核报告为 Word 文档。

    排版规范参照附注模板：
    - 页边距：左3cm、右3.18cm、上3.2cm、下2.54cm
    - 中文字体：仿宋_GB2312，小四号(12pt)；表格内五号(10.5pt)
    - 英文/数字字体：Arial Narrow
    - 段落间距：段前0行、段后0.9行，单倍行距
    - 表格：上下边框1磅，标题行下边框0.5磅，标题行及合计行加粗，
      首列左对齐，其余列居中，垂直居中
    """
    import io
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from fastapi.responses import Response

    # ── 排版常量 ──
    CN_FONT = "仿宋_GB2312"
    EN_FONT = "Arial Narrow"
    BODY_SIZE = Pt(12)       # 小四号
    TABLE_SIZE = Pt(10.5)    # 五号
    SMALL_SIZE = Pt(9)       # 小五号

    def _set_run_font(run, size=BODY_SIZE, bold=False, color=None):
        """统一设置 run 的中英文字体、字号、加粗、颜色。"""
        run.font.name = EN_FONT
        run.font.size = size
        run.bold = bold
        r = run._element
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), CN_FONT)
        rfonts.set(qn("w:ascii"), EN_FONT)
        rfonts.set(qn("w:hAnsi"), EN_FONT)
        if color:
            run.font.color.rgb = color

    def _set_para_spacing(para, before=0, after=0.9, line=1.0):
        """设置段落间距（单位：行）。"""
        fmt = para.paragraph_format
        fmt.space_before = Pt(before * 12)
        fmt.space_after = Pt(after * 12)
        fmt.line_spacing = line

    def _add_body_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                       before=0, after=0.9):
        """添加正文段落，自动应用字体和间距。"""
        p = doc.add_paragraph()
        p.alignment = align
        _set_para_spacing(p, before=before, after=after)
        run = p.add_run(text)
        _set_run_font(run, size=BODY_SIZE, bold=bold)
        return p

    def _set_table_borders(table):
        """设置表格边框：上下1磅，内部0.5磅，左右无。"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
        # 移除已有的 tblBorders
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tblPr.append(borders)

    def _set_header_bottom_border(row):
        """给表头行下方设置 0.5 磅边框。"""
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.tcPr if tc.tcPr is not None else parse_xml(
                f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                '</w:tcBorders>'
            )
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcPr.append(borders)
            if tc.tcPr is None:
                tc.append(tcPr)

    def _format_table_cell(cell, text, size=TABLE_SIZE, bold=False,
                           align=WD_ALIGN_PARAGRAPH.CENTER):
        """格式化表格单元格：设置文本、字体、对齐、垂直居中。"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        _set_para_spacing(p, before=0, after=0, line=1.0)
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    try:
        session_id = req.session_id

        findings = _findings.get(session_id, [])
        confirmed = [f for f in findings
                     if f.confirmation_status == FindingConfirmationStatus.CONFIRMED]
        output = (confirmed if confirmed
                  else [f for f in findings
                        if f.confirmation_status != FindingConfirmationStatus.DISMISSED])

        doc = Document()

        # ── 页面设置：页边距 ──
        for section in doc.sections:
            section.left_margin = Cm(3)
            section.right_margin = Cm(3.18)
            section.top_margin = Cm(3.2)
            section.bottom_margin = Cm(2.54)
            section.header_distance = Cm(1.3)
            section.footer_distance = Cm(1.3)

        # ── 设置默认样式字体 ──
        style = doc.styles["Normal"]
        style.font.name = EN_FONT
        style.font.size = BODY_SIZE
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

        # ── 文档标题 ──
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(title_p, before=1.0, after=0.5)
        title_run = title_p.add_run("审计报告复核结果")
        _set_run_font(title_run, size=Pt(18), bold=True)

        # ── 基本信息 ──
        _add_body_para(doc, f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
                       before=0.5, after=0.3)
        _add_body_para(doc, f"问题总数：{len(output)} 项", after=0.3)

        # ── 汇总统计表 ──
        risk_map = {"high": "高", "medium": "中", "low": "低"}
        cat_map = {
            "amount_inconsistency": "金额不一致",
            "reconciliation_error": "勾稽错误",
            "change_abnormal": "变动异常",
            "note_missing": "附注缺失",
            "report_body_compliance": "正文规范性",
            "note_content": "附注内容",
            "text_quality": "文本质量",
            "manual_annotation": "复核批注",
            "expression_compliance": "表达合规性",
            "missing_disclosure": "披露缺失",
            "other": "其他",
        }

        if output:
            # 按风险等级统计
            risk_counts: Dict[str, int] = {"high": 0, "medium": 0, "low": 0}
            for f in output:
                rk = f.risk_level.value if hasattr(f.risk_level, "value") else str(f.risk_level)
                risk_counts[rk] = risk_counts.get(rk, 0) + 1

            _add_body_para(doc, "一、问题汇总", bold=True, before=0.5, after=0.5)

            # 风险汇总表
            summary_tbl = doc.add_table(rows=2, cols=4)
            summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(summary_tbl)

            for i, hdr_text in enumerate(["风险等级", "高风险", "中风险", "低风险"]):
                _format_table_cell(summary_tbl.rows[0].cells[i], hdr_text,
                                   bold=True, size=TABLE_SIZE)
            _set_header_bottom_border(summary_tbl.rows[0])

            _format_table_cell(summary_tbl.rows[1].cells[0], "数量",
                               bold=False, size=TABLE_SIZE)
            for j, rk in enumerate(["high", "medium", "low"]):
                _format_table_cell(summary_tbl.rows[1].cells[j + 1],
                                   str(risk_counts.get(rk, 0)),
                                   size=TABLE_SIZE)

            # 表格后间距
            _add_body_para(doc, "", before=0.5, after=0.3)

            # ── 按分类分组的问题明细 ──
            _add_body_para(doc, "二、问题明细", bold=True, before=0.5, after=0.5)

            # 分类顺序（与前端一致）
            CATEGORY_ORDER = [
                "manual_annotation",
                "amount_inconsistency", "reconciliation_error", "change_abnormal",
                "note_missing", "note_content",
                "report_body_compliance", "text_quality",
            ]
            cat_grouped: Dict[str, list] = {}
            cat_order: list = []
            for cat_key in CATEGORY_ORDER:
                items = [f for f in output if (f.category.value if hasattr(f.category, "value") else str(f.category)) == cat_key]
                if not items:
                    continue
                label = cat_map.get(cat_key, cat_key)
                cat_grouped[label] = items
                cat_order.append(label)
            # 兜底：未在 CATEGORY_ORDER 中的分类
            for f in output:
                cat_val = f.category.value if hasattr(f.category, "value") else str(f.category)
                if cat_val not in CATEGORY_ORDER:
                    label = cat_map.get(cat_val, cat_val)
                    if label not in cat_grouped:
                        cat_grouped[label] = []
                        cat_order.append(label)
                    cat_grouped[label].append(f)

            cat_idx = 0
            for cat_label in cat_order:
                flist = cat_grouped[cat_label]
                cat_idx += 1
                _add_body_para(doc, f"（{cat_idx}）{cat_label}（{len(flist)}项）",
                               bold=True, before=0.5, after=0.3)

                # 问题表格：5列（序号、科目、风险、位置、描述/建议）
                col_count = 5
                tbl = doc.add_table(rows=1, cols=col_count)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                _set_table_borders(tbl)

                col_widths = [Cm(0.8), Cm(2.5), Cm(0.8), Cm(3.2), Cm(7.8)]
                for ci, w in enumerate(col_widths):
                    tbl.columns[ci].width = w

                hdr_texts = ["序号", "科目", "风险", "位置", "描述及建议"]
                for ci, ht in enumerate(hdr_texts):
                    _format_table_cell(tbl.rows[0].cells[ci], ht,
                                       bold=True, size=TABLE_SIZE)
                _set_header_bottom_border(tbl.rows[0])

                # 数据行
                for fi, f in enumerate(flist):
                    row = tbl.add_row()
                    risk_val = f.risk_level.value if hasattr(f.risk_level, "value") else str(f.risk_level)

                    desc_parts = []
                    if f.description:
                        desc_parts.append(f.description)
                    if f.suggestion:
                        desc_parts.append(f"建议：{f.suggestion}")
                    desc_text = "\n".join(desc_parts)

                    cell_data = [
                        (str(fi + 1), WD_ALIGN_PARAGRAPH.CENTER),
                        (f.account_name or "", WD_ALIGN_PARAGRAPH.CENTER),
                        (risk_map.get(risk_val, risk_val), WD_ALIGN_PARAGRAPH.CENTER),
                        (f.location or "", WD_ALIGN_PARAGRAPH.LEFT),
                        (desc_text, WD_ALIGN_PARAGRAPH.LEFT),
                    ]
                    for ci, (text, align) in enumerate(cell_data):
                        _format_table_cell(row.cells[ci], text,
                                           size=TABLE_SIZE, align=align)

                    # 高风险行标红
                    if risk_val == "high":
                        for ci in range(col_count):
                            for p in row.cells[ci].paragraphs:
                                for run in p.runs:
                                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

                # 表格后空行
                _add_body_para(doc, "", before=0.5, after=0.3)

        else:
            _add_body_para(doc, "未发现需要关注的问题。", before=0.5)

        # ── 页脚：页码 ──
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(fp, before=0, after=0)
            # 插入页码域代码
            run = fp.add_run()
            fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._element.append(fld_char_begin)
            run2 = fp.add_run()
            instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run2._element.append(instr)
            run3 = fp.add_run()
            fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3._element.append(fld_char_end)
            for r in [run, run2, run3]:
                _set_run_font(r, size=SMALL_SIZE)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f'audit_review_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        display_name = f'审计报告复核结果_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=\"{filename}\"; filename*=UTF-8''{requests_quote(display_name)}",
            },
        )
    except Exception as e:
        logger.exception("导出复核报告失败")
        raise HTTPException(500, f"导出失败: {str(e)}")


@router.patch("/finding/{finding_id}/status")
async def update_finding_status(finding_id: str, req: StatusUpdateRequest):
    """更新已确认问题的处理状态。"""
    finding = _find_finding(finding_id)
    if not finding:
        raise HTTPException(404, "Finding 不存在")
    finding.status = FindingStatus(req.status)
    return {"status": req.status}


# ─── 源文档预览 ───

@router.get("/source-preview/{file_id}")
async def source_preview(file_id: str):
    """获取源文档预览数据。"""
    return {"file_id": file_id, "content_html": "<p>预览功能待实现</p>"}


@router.get("/source-preview/{file_id}/sheet/{sheet_name}")
async def source_preview_sheet(file_id: str, sheet_name: str):
    """获取指定 Sheet 的预览数据。"""
    return {"file_id": file_id, "sheet_name": sheet_name, "content_html": "<p>Sheet预览待实现</p>"}


# ─── 母公司附注预设科目 ───

@router.get("/parent-accounts/{template_type}")
async def get_parent_accounts(template_type: str):
    """获取母公司附注预设科目清单。"""
    from ..services.account_mapping_template import account_mapping_template
    accounts = account_mapping_template.get_parent_company_accounts(template_type)
    if not accounts:
        raise HTTPException(400, f"不支持的模板类型: {template_type}")
    return {"template_type": template_type, "accounts": accounts}


# ─── 合并附注预设科目 ───

@router.get("/consolidated-accounts/{template_type}")
async def get_consolidated_accounts(template_type: str):
    """获取合并附注预设科目清单。"""
    from ..services.account_mapping_template import account_mapping_template
    accounts = account_mapping_template.get_consolidated_accounts(template_type)
    if not accounts:
        raise HTTPException(400, f"不支持的模板类型: {template_type}")
    return {"template_type": template_type, "accounts": accounts}


# ─── 模板管理 ───

@router.get("/templates/{template_type}/{template_category}")
async def get_template(template_type: str, template_category: str):
    """获取模板内容。"""
    try:
        tt = ReportTemplateType(template_type)
        tc = TemplateCategory(template_category)
    except ValueError as e:
        raise HTTPException(400, str(e))

    doc = report_template_service.get_template(tt, tc)
    if not doc:
        raise HTTPException(404, "模板未找到")
    return doc.model_dump()


@router.get("/templates/{template_type}/{template_category}/toc")
async def get_template_toc(template_type: str, template_category: str):
    """获取模板目录结构。"""
    try:
        tt = ReportTemplateType(template_type)
        tc = TemplateCategory(template_category)
    except ValueError as e:
        raise HTTPException(400, str(e))

    toc = report_template_service.get_template_toc(tt, tc)
    return {"toc": [e.model_dump() for e in toc]}


@router.get("/templates/{template_type}/{template_category}/section")
async def get_template_section(template_type: str, template_category: str, path: str = ""):
    """按路径获取模板章节内容。"""
    try:
        tt = ReportTemplateType(template_type)
        tc = TemplateCategory(template_category)
    except ValueError as e:
        raise HTTPException(400, str(e))

    content = report_template_service.get_template_section(tt, tc, path)
    if content is None:
        raise HTTPException(404, "章节未找到")
    return {"path": path, "content": content}


class UpdateTemplateRequest(BaseModel):
    content: str


@router.put("/templates/{template_type}/{template_category}")
async def update_template(template_type: str, template_category: str, req: UpdateTemplateRequest):
    """更新模板内容。"""
    try:
        tt = ReportTemplateType(template_type)
        tc = TemplateCategory(template_category)
    except ValueError as e:
        raise HTTPException(400, str(e))

    try:
        doc = report_template_service.update_template(tt, tc, req.content)
        return doc.model_dump()
    except ValueError as e:
        raise HTTPException(422, str(e))


@router.post("/templates/import")
async def import_template(
    file: UploadFile = File(...),
    template_type: str = Form(...),
    template_category: str = Form(...),
):
    """从 Word 导入模板。"""
    try:
        tt = ReportTemplateType(template_type)
        tc = TemplateCategory(template_category)
    except ValueError as e:
        raise HTTPException(400, str(e))

    content = await file.read()
    try:
        doc = report_template_service.import_from_word(content, tt, tc)
        return doc.model_dump()
    except ValueError as e:
        raise HTTPException(422, str(e))


# ─── Helper ───

def _find_finding(finding_id: str) -> Optional[ReportReviewFinding]:
    for session_findings in _findings.values():
        for f in session_findings:
            if f.id == finding_id:
                return f
    return None


def _find_finding_session_id(finding_id: str) -> Optional[str]:
    """找到 finding 所属的 session_id。"""
    for sid, session_findings in _findings.items():
        for f in session_findings:
            if f.id == finding_id:
                return sid
    return None


def _persist_finding_change(finding_id: str) -> None:
    """finding 变更后持久化对应 session 的 findings。"""
    sid = _find_finding_session_id(finding_id)
    if sid and sid in _findings:
        session_store.save_findings(sid, _findings[sid])
